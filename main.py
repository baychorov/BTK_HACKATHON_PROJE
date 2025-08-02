import streamlit as st
import google.generativeai as genai
import sqlite3
from dotenv import load_dotenv
import os
import json
from io import BytesIO
from datetime import datetime

# ReportLab - Türkçe karakterler için en iyi seçenek
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import platform
    import os

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    st.warning("ReportLab kütüphanesi bulunamadı. PDF indirme için: pip install reportlab")

# Python-docx for Word files
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("python-docx kütüphanesi bulunamadı. Word indirme için: pip install python-docx")

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=api_key)

model = genai.GenerativeModel("gemini-2.5-flash")

# Veritabanı bağlantısı - Yeni yapı
conn = sqlite3.connect("historai.db", check_same_thread=False)
c = conn.cursor()

# Yeni tablo yapısı: conversations (sohbetler) ve messages (mesajlar)
c.execute('''CREATE TABLE IF NOT EXISTS conversations 
             (id INTEGER PRIMARY KEY AUTOINCREMENT, 
              character TEXT, 
              title TEXT,
              created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
              is_pinned BOOLEAN DEFAULT 0)''')

c.execute('''CREATE TABLE IF NOT EXISTS messages 
             (id INTEGER PRIMARY KEY AUTOINCREMENT, 
              conversation_id INTEGER, 
              question TEXT, 
              answer TEXT, 
              created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
              FOREIGN KEY (conversation_id) REFERENCES conversations (id))''')

# Eski tablo yapısından geçiş
c.execute("PRAGMA table_info(chats)")
old_table_exists = c.fetchall()
if old_table_exists:
    # Eski verileri yeni yapıya taşı
    c.execute("SELECT character, question, answer FROM chats")
    old_chats = c.fetchall()
    for char, ques, ans in old_chats:
        # Her eski sohbet için yeni conversation oluştur
        c.execute("INSERT INTO conversations (character, title) VALUES (?, ?)",
                  (char, ques[:50] + "..." if len(ques) > 50 else ques))
        conv_id = c.lastrowid
        c.execute("INSERT INTO messages (conversation_id, question, answer) VALUES (?, ?, ?)",
                  (conv_id, ques, ans))
    # Eski tabloyu sil
    c.execute("DROP TABLE chats")

conn.commit()

st.set_page_config(page_title="HistorAI", layout="wide")
st.title("🧙‍♂ HistorAI - Tarihi Karakter Chatbotu")

# Session state başlatma
if "current_conversation_id" not in st.session_state:
    st.session_state.current_conversation_id = None
if "current_character" not in st.session_state:
    st.session_state.current_character = ""
if "messages" not in st.session_state:
    st.session_state.messages = []

# Sağ panel: Sohbet geçmişi
with st.sidebar:
    st.header("📚 Geçmiş Sohbetler")

    # Filtreleme
    filter_char = st.text_input("Karaktere göre filtrele")

    # Sohbetleri getir (sabitlenenler önce)
    if filter_char:
        c.execute("""SELECT id, character, title, is_pinned FROM conversations 
                    WHERE character LIKE ? ORDER BY is_pinned DESC, created_at DESC""",
                  ('%' + filter_char + '%',))
    else:
        c.execute("""SELECT id, character, title, is_pinned FROM conversations 
                    ORDER BY is_pinned DESC, created_at DESC""")
    conversations = c.fetchall()

    # Sohbet listesi
    for conv_id, char, title, is_pinned in conversations:
        pin_icon = "📌 " if is_pinned else ""
        label = f"{pin_icon}{char}: {title[:25]}..."

        col1, col2, col3 = st.columns([4, 1, 1])
        with col1:
            if st.button(label, key=f"conv_{conv_id}"):
                st.session_state.current_conversation_id = conv_id
                st.session_state.current_character = char
                # Mevcut sohbetin mesajlarını yükle
                c.execute("SELECT question, answer FROM messages WHERE conversation_id = ? ORDER BY created_at",
                          (conv_id,))
                messages = c.fetchall()
                st.session_state.messages = []
                for q, a in messages:
                    st.session_state.messages.append({"role": "user", "content": q})
                    st.session_state.messages.append({"role": "assistant", "content": a})
                st.rerun()

        with col2:
            # Pin/Unpin butonu
            pin_text = "📌" if not is_pinned else "📍"
            if st.button(pin_text, key=f"pin_{conv_id}"):
                new_pin_status = 0 if is_pinned else 1
                c.execute("UPDATE conversations SET is_pinned = ? WHERE id = ?",
                          (new_pin_status, conv_id))
                conn.commit()
                st.rerun()

        with col3:
            # Sil butonu
            if st.button("🗑", key=f"del_{conv_id}"):
                c.execute("DELETE FROM messages WHERE conversation_id = ?", (conv_id,))
                c.execute("DELETE FROM conversations WHERE id = ?", (conv_id,))
                conn.commit()
                if st.session_state.current_conversation_id == conv_id:
                    st.session_state.current_conversation_id = None
                    st.session_state.messages = []
                st.rerun()

    st.divider()

    # Yeni sohbet başlat
    if st.button("✨ Yeni Sohbet Başlat"):
        st.session_state.current_conversation_id = None
        st.session_state.current_character = ""
        st.session_state.messages = []
        st.rerun()

    # Tüm geçmişi sil
    if st.button("🧨 Tüm Geçmişi Sil"):
        c.execute("DELETE FROM messages")
        c.execute("DELETE FROM conversations")
        conn.commit()
        st.session_state.current_conversation_id = None
        st.session_state.messages = []
        st.rerun()

    st.divider()

    # İndirme seçenekleri
    st.subheader("📥 İndirme Seçenekleri")

    if st.session_state.current_conversation_id:
        st.write("Mevcut sohbeti indir:")


        # PDF indirme fonksiyonu
        def create_conversation_pdf(conversation_id):
            if not REPORTLAB_AVAILABLE:
                st.error("PDF oluşturmak için ReportLab gerekli: pip install reportlab")
                return None

            # Sohbet bilgilerini al
            c.execute("SELECT character, title FROM conversations WHERE id = ?", (conversation_id,))
            conv_info = c.fetchone()
            if not conv_info:
                return None

            character, title = conv_info
            c.execute("SELECT question, answer FROM messages WHERE conversation_id = ? ORDER BY created_at",
                      (conversation_id,))
            messages = c.fetchall()

            def register_modern_fonts():
                """Modern font kaydı"""
                try:
                    if platform.system() == "Windows":
                        font_paths = [
                            "C:/Windows/Fonts/segoeui.ttf",  # Segoe UI - modern
                            "C:/Windows/Fonts/calibri.ttf",
                            "C:/Windows/Fonts/arial.ttf"
                        ]
                        bold_paths = [
                            "C:/Windows/Fonts/segoeuib.ttf",
                            "C:/Windows/Fonts/calibrib.ttf",
                            "C:/Windows/Fonts/arialbd.ttf"
                        ]
                    elif platform.system() == "Darwin":
                        font_paths = [
                            "/System/Library/Fonts/SF-Pro-Text-Regular.otf",
                            "/System/Library/Fonts/Helvetica.ttc",
                            "/Library/Fonts/Arial.ttf"
                        ]
                        bold_paths = [
                            "/System/Library/Fonts/SF-Pro-Text-Bold.otf",
                            "/System/Library/Fonts/Helvetica-Bold.ttc",
                            "/Library/Fonts/Arial Bold.ttf"
                        ]
                    else:
                        font_paths = [
                            "/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf",
                            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
                        ]
                        bold_paths = [
                            "/usr/share/fonts/truetype/ubuntu/Ubuntu-B.ttf",
                            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
                        ]

                    # Font kaydı
                    for i, (regular, bold) in enumerate(zip(font_paths, bold_paths)):
                        try:
                            if os.path.exists(regular) and os.path.exists(bold):
                                pdfmetrics.registerFont(TTFont('ModernFont', regular))
                                pdfmetrics.registerFont(TTFont('ModernFont-Bold', bold))
                                return 'ModernFont'
                        except:
                            continue

                    return 'Helvetica'
                except:
                    return 'Helvetica'

            modern_font = register_modern_fonts()
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50,
                                    topMargin=50, bottomMargin=50)

            styles = getSampleStyleSheet()

            # Modern stiller
            title_style = ParagraphStyle(
                'ModernTitle',
                parent=styles['Heading1'],
                fontSize=24,
                alignment=1,
                spaceAfter=30,
                textColor=HexColor('#1a365d'),
                fontName=f'{modern_font}-Bold' if modern_font != 'Helvetica' else 'Helvetica-Bold'
            )

            subtitle_style = ParagraphStyle(
                'ModernSubtitle',
                parent=styles['Heading2'],
                fontSize=16,
                alignment=1,
                spaceAfter=20,
                textColor=HexColor('#2d3748'),
                fontName=f'{modern_font}-Bold' if modern_font != 'Helvetica' else 'Helvetica-Bold'
            )

            question_style = ParagraphStyle(
                'ModernQuestion',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=8,
                spaceBefore=15,
                textColor=HexColor('#2b6cb0'),
                fontName=f'{modern_font}-Bold' if modern_font != 'Helvetica' else 'Helvetica-Bold'
            )

            answer_style = ParagraphStyle(
                'ModernAnswer',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=15,
                leading=16,
                textColor=HexColor('#1a202c'),
                fontName=modern_font
            )

            content = []
            content.append(Paragraph("🧙‍♂ HistorAI Sohbeti", title_style))
            content.append(Paragraph(f"Karakter: {character}", subtitle_style))
            content.append(Spacer(1, 30))

            for i, (question, answer) in enumerate(messages, 1):
                def clean_text(text):
                    return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                content.append(Paragraph(f"<b>❓ Soru {i}:</b> {clean_text(question)}", question_style))
                content.append(Paragraph(f"<b>💬 {character}:</b> {clean_text(answer)}", answer_style))

                if i < len(messages):
                    content.append(Spacer(1, 10))

            try:
                doc.build(content)
                buffer.seek(0)
                return buffer
            except Exception as e:
                st.error(f"PDF oluştururken hata: {str(e)}")
                return None


        # Word dosyası oluşturma
        def create_conversation_word(conversation_id):
            if not DOCX_AVAILABLE:
                st.error("Word dosyası oluşturmak için python-docx gerekli: pip install python-docx")
                return None

            c.execute("SELECT character, title FROM conversations WHERE id = ?", (conversation_id,))
            conv_info = c.fetchone()
            if not conv_info:
                return None

            character, title = conv_info
            c.execute("SELECT question, answer FROM messages WHERE conversation_id = ? ORDER BY created_at",
                      (conversation_id,))
            messages = c.fetchall()

            doc = Document()

            # Başlık
            title_para = doc.add_heading('🧙‍♂ HistorAI Sohbeti', 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Karakter bilgisi
            char_para = doc.add_heading(f'Karakter: {character}', level=1)
            char_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph('')

            # Mesajlar
            for i, (question, answer) in enumerate(messages, 1):
                # Soru
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f'❓ Soru {i}: ')
                q_run.bold = True
                q_run.font.color.rgb = RGBColor(43, 108, 176)
                q_para.add_run(question)

                # Cevap
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f'💬 {character}: ')
                a_run.bold = True
                a_run.font.color.rgb = RGBColor(212, 84, 58)
                a_para.add_run(answer)

                doc.add_paragraph('')

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer


        # JSON oluşturma
        def create_conversation_json(conversation_id):
            c.execute("SELECT character, title FROM conversations WHERE id = ?", (conversation_id,))
            conv_info = c.fetchone()
            if not conv_info:
                return None

            character, title = conv_info
            c.execute("SELECT question, answer FROM messages WHERE conversation_id = ? ORDER BY created_at",
                      (conversation_id,))
            messages = c.fetchall()

            data = {
                "character": character,
                "title": title,
                "messages": [{"question": q, "answer": a} for q, a in messages]
            }

            return BytesIO(json.dumps(data, indent=4, ensure_ascii=False).encode("utf-8"))


        # İndirme butonları
        col1, col2, col3 = st.columns(3)

        with col1:
            pdf_data = create_conversation_pdf(st.session_state.current_conversation_id)
            if pdf_data:
                st.download_button("📄 PDF", data=pdf_data,
                                   file_name=f"historai_{st.session_state.current_character}.pdf",
                                   mime="application/pdf")

        with col2:
            word_data = create_conversation_word(st.session_state.current_conversation_id)
            if word_data:
                st.download_button("📝 Word", data=word_data,
                                   file_name=f"historai_{st.session_state.current_character}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with col3:
            json_data = create_conversation_json(st.session_state.current_conversation_id)
            if json_data:
                st.download_button("🗂 JSON", data=json_data,
                                   file_name=f"historai_{st.session_state.current_character}.json",
                                   mime="application/json")

# Kişilik testi veri yapısı
PERSONALITY_TEST = {
    "questions": [
        {
            "question": "Yeni bir proje başlatırken hangi yaklaşımı tercih edersiniz?",
            "options": [
                {"text": "Detaylı plan yapar, her aşamayı önceden hesaplarım",
                 "traits": {"conscientiousness": 2, "openness": 1}},
                {"text": "Genel bir fikir ile başlar, yol boyunca şekillendiriririm",
                 "traits": {"openness": 2, "conscientiousness": -1}},
                {"text": "Başkalarının fikirlerini dinler, ortak karar veririm",
                 "traits": {"agreeableness": 2, "extraversion": 1}},
                {"text": "İçgüdülerime güvenir, spontane hareket ederim", "traits": {"neuroticism": 1, "openness": 1}}
            ]
        },
        {
            "question": "Karşılaştığınız zorluklar karşısında nasıl tepki verirsiniz?",
            "options": [
                {"text": "Analitik düşünür, sistematik çözümler ararım",
                 "traits": {"conscientiousness": 2, "neuroticism": -1}},
                {"text": "Yaratıcı ve sıra dışı yöntemler denerim", "traits": {"openness": 2, "conscientiousness": -1}},
                {"text": "Diğer insanlardan yardım ve tavsiye alırım",
                 "traits": {"agreeableness": 2, "extraversion": 1}},
                {"text": "Duygusal yaklaşır, içsel motivasyonuma güvenirim",
                 "traits": {"neuroticism": 1, "agreeableness": 1}}
            ]
        },
        {
            "question": "İdeal bir akşam nasıl geçirirsiniz?",
            "options": [
                {"text": "Kitap okuyarak veya öğrendiğim konuları derinleştirerek",
                 "traits": {"openness": 2, "extraversion": -1}},
                {"text": "Arkadaşlarımla sohbet ederek, deneyimlerimi paylaşarak",
                 "traits": {"extraversion": 2, "agreeableness": 1}},
                {"text": "Sanat, müzik veya yaratıcı aktivitelerle",
                 "traits": {"openness": 2, "conscientiousness": -1}},
                {"text": "Düzenli rutinlerimi sürdürerek, planlarımı gözden geçirerek",
                 "traits": {"conscientiousness": 2, "extraversion": -1}}
            ]
        },
        {
            "question": "Liderlik tarzınızı nasıl tanımlarsınız?",
            "options": [
                {"text": "Vizyon sahibi, ilham verici ve yenilikçi", "traits": {"openness": 2, "extraversion": 1}},
                {"text": "Disiplinli, adaletli ve kurallara bağlı",
                 "traits": {"conscientiousness": 2, "agreeableness": 1}},
                {"text": "Empati kuran, destekleyici ve işbirlikçi", "traits": {"agreeableness": 2, "extraversion": 1}},
                {"text": "Kararlı, tutarlı ancak esnek", "traits": {"conscientiousness": 1, "neuroticism": -1}}
            ]
        },
        {
            "question": "Hangi tür bilgi sizi en çok cezbeder?",
            "options": [
                {"text": "Bilimsel keşifler ve teknolojik yenilikler",
                 "traits": {"openness": 2, "conscientiousness": 1}},
                {"text": "Tarihsel olaylar ve kültürel gelişmeler", "traits": {"openness": 1, "conscientiousness": 1}},
                {"text": "İnsan ilişkileri ve sosyal dinamikler", "traits": {"agreeableness": 2, "extraversion": 1}},
                {"text": "Felsefe ve yaşamın anlamı üzerine düşünceler", "traits": {"openness": 2, "neuroticism": 1}}
            ]
        }
    ],
    "characters": [
        {
            "name": "Leonardo da Vinci",
            "traits": {"openness": 10, "conscientiousness": 7, "extraversion": 5, "agreeableness": 6, "neuroticism": 4},
            "description": "Çok yönlü deha, sanat ve bilimi birleştiren yaratıcı vizyon",
            "quote": "Öğrenme bizim yaşadığımız sürece devam eder."
        },
        {
            "name": "Fatih Sultan Mehmet",
            "traits": {"openness": 8, "conscientiousness": 9, "extraversion": 8, "agreeableness": 5, "neuroticism": 3},
            "description": "Stratejik düşünür, kararlı lider ve vizyon sahibi fatih",
            "quote": "Ya İstanbul'u alırım, ya da İstanbul beni alır."
        },
        {
            "name": "Mevlana",
            "traits": {"openness": 9, "conscientiousness": 6, "extraversion": 6, "agreeableness": 10, "neuroticism": 2},
            "description": "Sevgi dolu, hoşgörülü ve hakikati arayan mutasavvıf",
            "quote": "Sevgi yolculuğu, bizi kendimize götürür."
        },
        {
            "name": "Ibn Khaldun",
            "traits": {"openness": 9, "conscientiousness": 8, "extraversion": 4, "agreeableness": 7, "neuroticism": 3},
            "description": "Sosyal bilimcı, tarihçi ve medeniyet analisti",
            "quote": "Tarih, toplumların yükseliş ve çöküş kanunlarını öğretir."
        },
        {
            "name": "Yunus Emre",
            "traits": {"openness": 8, "conscientiousness": 5, "extraversion": 7, "agreeableness": 9, "neuroticism": 4},
            "description": "Halkın ozanı, sevgi ve kardeşlik şairi",
            "quote": "Yaratılanı severiz, Yaratan'dan ötürü."
        },
        {
            "name": "Sultan Alparslan",
            "traits": {"openness": 6, "conscientiousness": 9, "extraversion": 7, "agreeableness": 6, "neuroticism": 2},
            "description": "Adil hükümdar, stratejik komutan ve devlet adamı",
            "quote": "Adalet, saltanatın temelidir."
        },
        {
            "name": "İbn Sina (Avicenna)",
            "traits": {"openness": 10, "conscientiousness": 8, "extraversion": 4, "agreeableness": 7, "neuroticism": 3},
            "description": "Hekim, filozof ve bilim insanı",
            "quote": "Bilgi, onu arayan ve emek verenlerindir."
        },
        {
            "name": "Akşemseddin",
            "traits": {"openness": 8, "conscientiousness": 8, "extraversion": 5, "agreeableness": 8, "neuroticism": 2},
            "description": "Bilgin, mutasavvıf ve Fatih'in hocası",
            "quote": "İlim öğren, kendini bil."
        }
    ]
}


# Test sonucu hesaplama fonksiyonu
def calculate_personality_match(user_scores):
    """Kullanıcının kişilik puanlarını karakterlerle eşleştir"""
    best_matches = []

    for character in PERSONALITY_TEST["characters"]:
        # Her karakter için benzerlik puanı hesapla
        similarity_score = 0
        total_possible = 0

        for trait in ["openness", "conscientiousness", "extraversion", "agreeableness", "neuroticism"]:
            user_score = user_scores.get(trait, 0)
            char_score = character["traits"][trait]

            # Mutlak farkı hesapla (0-10 arası)
            diff = abs(user_score - char_score)
            similarity = 10 - diff  # Fark ne kadar az ise benzerlik o kadar yüksek

            similarity_score += similarity
            total_possible += 10

        # Yüzde olarak hesapla
        match_percentage = (similarity_score / total_possible) * 100

        best_matches.append({
            "character": character,
            "percentage": match_percentage
        })

    # En yüksek eşleşenleri döndür
    return sorted(best_matches, key=lambda x: x["percentage"], reverse=True)


# Session state için test değişkenleri
if "test_active" not in st.session_state:
    st.session_state.test_active = False
if "test_question_index" not in st.session_state:
    st.session_state.test_question_index = 0
if "test_scores" not in st.session_state:
    st.session_state.test_scores = {"openness": 0, "conscientiousness": 0, "extraversion": 0, "agreeableness": 0,
                                    "neuroticism": 0}
if "test_completed" not in st.session_state:
    st.session_state.test_completed = False

# Ana sohbet alanı
if not st.session_state.current_conversation_id:
    # Tarihî Kişilik Testi Bölümü
    st.markdown("---")

    col1, col2 = st.columns([2, 1])

    with col1:
        st.markdown("### 🧬 Tarihî Kişilik Testi")
        st.markdown("**5 soruluk kısa testle hangi tarihi karaktere benzediğinizi keşfedin!**")

        if not st.session_state.test_active and not st.session_state.test_completed:
            if st.button("🚀 Teste Başla", type="primary"):
                st.session_state.test_active = True
                st.session_state.test_question_index = 0
                st.session_state.test_scores = {"openness": 0, "conscientiousness": 0, "extraversion": 0,
                                                "agreeableness": 0, "neuroticism": 0}
                st.rerun()

    with col2:
        st.markdown("#### 🎯 Test Sonrası:")
        st.markdown("✨ Kişilik eşleşmesi  \n📊 Uyumluluk yüzdesi  \n💬 Direkt sohbet başlat")

    # Test aktifse soruları göster
    if st.session_state.test_active:
        current_q = st.session_state.test_question_index
        total_q = len(PERSONALITY_TEST["questions"])

        if current_q < total_q:
            st.markdown("---")
            # Progress bar
            progress = (current_q) / total_q
            st.progress(progress, text=f"Soru {current_q + 1} / {total_q}")

            question_data = PERSONALITY_TEST["questions"][current_q]

            st.markdown(f"### 📝 Soru {current_q + 1}")
            st.markdown(f"**{question_data['question']}**")

            # Seçenekleri radio button olarak göster
            option_labels = [opt["text"] for opt in question_data["options"]]

            selected_option = st.radio(
                "Seçiminizi yapın:",
                options=range(len(option_labels)),
                format_func=lambda x: option_labels[x],
                key=f"test_q_{current_q}"
            )

            col1, col2, col3 = st.columns([1, 1, 2])

            with col2:
                if st.button("➡️ Sonraki Soru", type="primary"):
                    # Seçilen seçeneğin trait puanlarını ekle
                    selected_traits = question_data["options"][selected_option]["traits"]
                    for trait, score in selected_traits.items():
                        st.session_state.test_scores[trait] += score

                    st.session_state.test_question_index += 1
                    st.rerun()

            with col1:
                if st.button("❌ Testi Durdur"):
                    st.session_state.test_active = False
                    st.session_state.test_question_index = 0
                    st.rerun()

        else:
            # Test tamamlandı - sonuçları göster
            st.session_state.test_active = False
            st.session_state.test_completed = True
            st.rerun()

    # Test sonuçları
    if st.session_state.test_completed:
        st.markdown("---")
        st.markdown("## 🎉 Test Sonuçlarınız")

        # Skorları normalize et (0-10 arası)
        normalized_scores = {}
        for trait, score in st.session_state.test_scores.items():
            # Her soru için maksimum 2 puan alınabilir, 5 soru var
            max_possible = 10  # 5 soru x 2 puan
            min_possible = -5  # 5 soru x -1 puan (bazı negatif skorlar var)

            # 0-10 arasına normalize et
            normalized = ((score - min_possible) / (max_possible - min_possible)) * 10
            normalized_scores[trait] = max(0, min(10, normalized))

        # En iyi eşleşmeleri bul
        matches = calculate_personality_match(normalized_scores)

        # En iyi 3 eşleşmeyi göster
        for i, match in enumerate(matches[:3]):
            character = match["character"]
            percentage = match["percentage"]

            if i == 0:
                # En iyi eşleşme - özel stil
                st.markdown(f"""
                <div style="
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    padding: 20px;
                    border-radius: 15px;
                    color: white;
                    margin: 10px 0;
                    text-align: center;
                ">
                    <h2>🏆 En İyi Eşleşmeniz!</h2>
                    <h1>{character['name']}</h1>
                    <h2>%{percentage:.0f} Uyumluluk</h2>
                    <p style="font-style: italic;">"{character['quote']}"</p>
                    <p>{character['description']}</p>
                </div>
                """, unsafe_allow_html=True)

                # Direkt sohbet başlatma butonu
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button(f"💬 {character['name']} ile Sohbet Başlat", type="primary", key="start_chat_best"):
                        st.session_state.current_character = character['name']
                        c.execute("INSERT INTO conversations (character, title) VALUES (?, ?)",
                                  (character['name'], f"{character['name']} ile kişilik testi sohbeti"))
                        st.session_state.current_conversation_id = c.lastrowid
                        conn.commit()
                        st.session_state.test_completed = False
                        st.rerun()

            else:
                # Diğer eşleşmeler
                with st.expander(f"#{i + 1} - {character['name']} (%{percentage:.0f} uyumluluk)"):
                    st.markdown(f"**{character['description']}**")
                    st.markdown(f"*\"{character['quote']}\"*")

                    if st.button(f"💬 {character['name']} ile Sohbet Başlat", key=f"start_chat_{i}"):
                        st.session_state.current_character = character['name']
                        c.execute("INSERT INTO conversations (character, title) VALUES (?, ?)",
                                  (character['name'], f"{character['name']} ile kişilik testi sohbeti"))
                        st.session_state.current_conversation_id = c.lastrowid
                        conn.commit()
                        st.session_state.test_completed = False
                        st.rerun()

        # Testi tekrar alma butonu
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("🔄 Testi Tekrar Al"):
                st.session_state.test_active = False
                st.session_state.test_completed = False
                st.session_state.test_question_index = 0
                st.rerun()

        with col2:
            if st.button("➡️ Manuel Karakter Seç"):
                st.session_state.test_completed = False
                st.rerun()

    # Manuel karakter seçimi (test yapılmadıysa veya manuel seçim isteniyorsa)
    if not st.session_state.test_active and not st.session_state.test_completed:
        st.markdown("---")
        st.markdown("### 🎭 Veya Manuel Karakter Seçin")
        character = st.text_input("Tarihi karakter adını girin:",
                                  placeholder="Örn: Fatih Sultan Mehmet, Leonardo da Vinci, Mevlana...")

        if character:
            st.session_state.current_character = character
            # Yeni conversation oluştur
            c.execute("INSERT INTO conversations (character, title) VALUES (?, ?)",
                      (character, f"{character} ile sohbet"))
            st.session_state.current_conversation_id = c.lastrowid
            conn.commit()
            st.rerun()

else:
    # Mevcut sohbet
    st.markdown(f"### 🗣 {st.session_state.current_character} ile sohbet ediyorsunuz")

    # Sohbet geçmişini göster
    for message in st.session_state.messages:
        if message["role"] == "user":
            with st.chat_message("user"):
                st.write(message["content"])
        else:
            with st.chat_message("assistant"):
                st.write(message["content"])

    # Yeni mesaj girişi
    if prompt := st.chat_input("Sorunuzu yazın..."):
        # Kullanıcı mesajını ekle
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)

        # AI yanıtı oluştur
        with st.chat_message("assistant"):
            with st.spinner("Yanıt oluşturuluyor..."):
                ai_prompt = f"""
Sen yalnızca tarihsel olarak belgelenmiş, gerçek ve yaşamış karakterlerin rolünü yapabilirsin...

*TEMEL KURALLAR:*
1. Yalnızca insanlık tarihinde yaşamış, güvenilir tarihsel kaynaklarda yer alan kişiliklerin yerine geçebilirsin.
2. Her yanıtın tarihsel olarak doğrulanabilir olmalı. Uydurma bilgi, tahmin ya da kurgu içerik üretmek kesinlikle yasaktır.

*ROL YAPMAYI REDDETMEN GEREKEN DURUMLAR:*
- Gerçek olmayan, hayali veya anlamsız karakterler (örneğin: "Merhaba", "Kral Ejder", "Mehmet", "RobotX")
- Tarihsel figür olmayan çağdaş kişiler (örneğin: Elon Musk, Donald Trump, Britney Spears, Ronaldo)
- Türkiye Cumhuriyeti tarafından hassas kabul edilen kişi ve içerikler (örneğin: terör örgütleri ve terör örgütü kurucuları, suçlu nitelikteki insanlar)
- Dini, tanrısal veya kutsal figürler (örneğin: Tanrı, Hz. Muhammed, İsa)
- Küfür, cinsellik, hakaret ve toplumsal olarak hassas konular

*Bu tür isteklerde:*
- Nazikçe isteği reddet
- Kısa açıklama yap: "Bu kişi/talep, rol yapabileceğim güvenilir tarihsel içeriklere uygun değildir."
- *Kesinlikle hiçbir şekilde rol yapma veya bu kişiler adına konuşma.*

*Tarihsel Uydurma Yasağı:*
Eğer kullanıcı sana gerçek bir tarihi olayla ilgisi olmayan bir hikâye, konuşma, anı ya da deneyim sorduysa:
- Uydurma cevap verme.
- "Bu olay/kaynak tarihsel olarak doğrulanmış değildir." diyerek açıklama yap.
- Ancak istenen olay gerçekte yaşanmışsa, tarihsel bilgiye dayalı şekilde cevap verebilirsin.

---

Şimdi {st.session_state.current_character} olarak konuşuyorsun. Aşağıdaki soruyu, bu karakterin tarihsel gerçeklerine ve dönemin diline sadık kalarak cevapla:

- *Soru:* {prompt}
"""
                try:
                    response = model.generate_content(ai_prompt)
                    answer = response.text
                    st.write(answer)

                    # Mesajları kaydet
                    st.session_state.messages.append({"role": "assistant", "content": answer})
                    c.execute("INSERT INTO messages (conversation_id, question, answer) VALUES (?, ?, ?)",
                              (st.session_state.current_conversation_id, prompt, answer))
                    conn.commit()

                except Exception as e:
                    error_msg = f"Bir hata oluştu: {str(e)}"
                    st.error(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})

    # Sohbeti bitir butonu
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("✅ Sohbeti Bitir"):
            # Conversation title'ını güncelle
            if st.session_state.messages:
                first_question = st.session_state.messages[0]["content"] if st.session_state.messages[0][
                                                                                "role"] == "user" else "Sohbet"
                title = first_question[:50] + "..." if len(first_question) > 50 else first_question
                c.execute("UPDATE conversations SET title = ? WHERE id = ?",
                          (title, st.session_state.current_conversation_id))
                conn.commit()

            st.session_state.current_conversation_id = None
            st.session_state.current_character = ""
            st.session_state.messages = []
            st.success("Sohbet tamamlandı ve geçmişe kaydedildi!")
            st.rerun()