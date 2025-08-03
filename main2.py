import streamlit as st
import google.generativeai as genai
import sqlite3
from dotenv import load_dotenv
import os
import json
from io import BytesIO
from datetime import datetime
import random

# ReportLab - TÃ¼rkÃ§e karakterler iÃ§in en iyi seÃ§enek
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import platform
    import os

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    st.warning("ReportLab kÃ¼tÃ¼phanesi bulunamadÄ±. PDF indirme iÃ§in: pip install reportlab")

# Python-docx for Word files
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("python-docx kÃ¼tÃ¼phanesi bulunamadÄ±. Word indirme iÃ§in: pip install python-docx")

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=api_key)

model = genai.GenerativeModel("gemini-2.5-flash")

# VeritabanÄ± baÄŸlantÄ±sÄ± - Yeni yapÄ±
conn = sqlite3.connect("historai.db", check_same_thread=False)
c = conn.cursor()

# Yeni tablo yapÄ±sÄ±: conversations (sohbetler) ve messages (mesajlar)
c.execute('''CREATE TABLE IF NOT EXISTS conversations 
             (id INTEGER PRIMARY KEY AUTOINCREMENT, 
              character TEXT, 
              title TEXT,
              created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
              is_pinned BOOLEAN DEFAULT 0,
              conversation_type TEXT DEFAULT 'normal')''')

c.execute('''CREATE TABLE IF NOT EXISTS messages 
             (id INTEGER PRIMARY KEY AUTOINCREMENT, 
              conversation_id INTEGER, 
              question TEXT, 
              answer TEXT, 
              created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
              FOREIGN KEY (conversation_id) REFERENCES conversations (id))''')

# Eski tablo yapÄ±sÄ±ndan geÃ§iÅŸ
c.execute("PRAGMA table_info(chats)")
old_table_exists = c.fetchall()
if old_table_exists:
    # Eski verileri yeni yapÄ±ya taÅŸÄ±
    c.execute("SELECT character, question, answer FROM chats")
    old_chats = c.fetchall()
    for char, ques, ans in old_chats:
        # Her eski sohbet iÃ§in yeni conversation oluÅŸtur
        c.execute("INSERT INTO conversations (character, title) VALUES (?, ?)",
                  (char, ques[:50] + "..." if len(ques) > 50 else ques))
        conv_id = c.lastrowid
        c.execute("INSERT INTO messages (conversation_id, question, answer) VALUES (?, ?, ?)",
                  (conv_id, ques, ans))
    # Eski tabloyu sil
    c.execute("DROP TABLE chats")

conn.commit()

st.set_page_config(page_title="HistorAI", layout="wide", page_icon="ğŸ§™â€â™‚")

# Session state baÅŸlatma
if "current_conversation_id" not in st.session_state:
    st.session_state.current_conversation_id = None
if "current_character" not in st.session_state:
    st.session_state.current_character = ""
if "messages" not in st.session_state:
    st.session_state.messages = []
if "current_page" not in st.session_state:
    st.session_state.current_page = "home"
if "time_travel_active" not in st.session_state:
    st.session_state.time_travel_active = False

# Tarihsel olaylar veri yapÄ±sÄ±
HISTORICAL_EVENTS = {
    "1453": {
        "event": "Ä°stanbul'un Fethi",
        "date": "29 MayÄ±s 1453",
        "characters": ["Fatih Sultan Mehmet", "Konstantin XI", "Halil PaÅŸa"],
        "setting": "Konstantinopolis surlarÄ± Ã¶nÃ¼nde top sesleri yankÄ±lanÄ±yor. Fatih Sultan Mehmet son hazÄ±rlÄ±klarÄ± gÃ¶zden geÃ§iriyor.",
        "opening": "Åafak vakti, sur dibindeki Ã§adÄ±rda Fatih Sultan Mehmet haritaya bakÄ±yor ve sana dÃ¶nerek diyor: 'Bu gece tarih yazÄ±lacak. Sen bu kutsal anda bizimle misin?'"
    },
    "1071": {
        "event": "Malazgirt SavaÅŸÄ±",
        "date": "26 AÄŸustos 1071",
        "characters": ["Sultan Alparslan", "Romanos IV", "Nizam-Ã¼l MÃ¼lk"],
        "setting": "Malazgirt ovalarÄ±nda iki bÃ¼yÃ¼k ordu karÅŸÄ± karÅŸÄ±ya. Sultan Alparslan beyaz kaftan giymiÅŸ, attan iniyor.",
        "opening": "Sultan Alparslan kÄ±lÄ±cÄ±nÄ± Ã§Ä±kararak topraÄŸa saplar ve sana bakar: 'EÄŸer ÅŸehit dÃ¼ÅŸersem, oÄŸlum MelikÅŸah'a bu kÄ±lÄ±cÄ± gÃ¶tÃ¼r. Sen bu tarihi anÄ±n ÅŸahidi olmaya hazÄ±r mÄ±sÄ±n?'"
    },
    "1789": {
        "event": "FransÄ±z Devrimi",
        "date": "14 Temmuz 1789",
        "characters": ["Robespierre", "Danton", "Marat"],
        "setting": "Paris sokaklarÄ±nda barikatlar kuruluyor, halk Bastille'e yÃ¼rÃ¼yor. Robespierre bir kahvehane kÃ¶ÅŸesinde planlar yapÄ±yor.",
        "opening": "Robespierre gÃ¶zlerinin iÃ§ine bakÄ±yor ve soruyor: 'Sen kimin tarafÄ±ndasÄ±n? KralÄ±n mÄ±, yoksa halkÄ±n mÄ±? Bu devrim iÃ§in kanÄ±n akacak!'"
    },
    "1492": {
        "event": "Amerika'nÄ±n KeÅŸfi",
        "date": "12 Ekim 1492",
        "characters": ["Kristof Kolomb", "Martin Pinzon", "Rodrigo de Triana"],
        "setting": "Santa Maria gemisinin gÃ¼vertesinde, uzun deniz yolculuÄŸunun ardÄ±ndan nihayet kara gÃ¶rÃ¼nÃ¼yor.",
        "opening": "Kolomb geminin direÄŸinde duruyor ve sana dÃ¶nerek diyor: 'Ä°ÅŸte! Yeni bir dÃ¼nya! Sen bu tarihi keÅŸfin tanÄ±ÄŸÄ± olmak ister misin?'"
    },
    "1299": {
        "event": "OsmanlÄ± Devleti'nin KuruluÅŸu",
        "date": "1299",
        "characters": ["Osman Gazi", "Åeyh Edebali", "Malhun Hatun"],
        "setting": "SÃ¶ÄŸÃ¼t'te kÃ¼Ã§Ã¼k bir beylik kurulmaya Ã§alÄ±ÅŸÄ±lÄ±yor. Osman Gazi Ã§adÄ±rÄ±nda gelecek planlarÄ± yapÄ±yor.",
        "opening": "Osman Gazi rÃ¼yasÄ±nÄ± anlatÄ±yor: 'GÃ¶ÄŸsÃ¼mden bir aÄŸaÃ§ Ã§Ä±ktÄ±, dallarÄ± tÃ¼m dÃ¼nyayÄ± kapladÄ±. Sen bu rÃ¼yanÄ±n gerÃ§ek olacaÄŸÄ±na inanÄ±r mÄ±sÄ±n?'"
    }
}


# Karakter tavsiye sistemi
def analyze_conversation_style(messages):
    """KonuÅŸma tarzÄ±nÄ± analiz ederek karakter Ã¶nerisi yapar"""
    if len(messages) < 4:  # En az 2 soru-cevap
        return None

    user_messages = [msg["content"].lower() for msg in messages if msg["role"] == "user"]
    all_text = " ".join(user_messages)

    suggestions = []

    # Felsefe ve dÃ¼ÅŸÃ¼nce aÄŸÄ±rlÄ±klÄ±
    philosophy_keywords = ["neden", "nasÄ±l", "anlam", "dÃ¼ÅŸÃ¼nce", "felsefe", "hakikat", "bilgi", "akÄ±l"]
    if any(keyword in all_text for keyword in philosophy_keywords):
        suggestions.extend([
            {"name": "Sokrates", "reason": "Felsefi sorgulamalarÄ±nÄ±z Sokrates'in tarzÄ±na Ã§ok benziyor"},
            {"name": "Ä°bn RÃ¼ÅŸd", "reason": "AkÄ±l ve mantÄ±k odaklÄ± yaklaÅŸÄ±mÄ±nÄ±z Ä°bn RÃ¼ÅŸd ile uyumlu"},
            {"name": "Farabi", "reason": "Bilgi arayÄ±ÅŸÄ±nÄ±z Farabi'nin yÃ¶ntemleriyle Ã¶rtÃ¼ÅŸÃ¼yor"}
        ])

    # SavaÅŸ ve strateji
    war_keywords = ["savaÅŸ", "strateji", "ordu", "zafer", "mÃ¼cadele", "liderlik"]
    if any(keyword in all_text for keyword in war_keywords):
        suggestions.extend([
            {"name": "Selahaddin Eyyubi", "reason": "Strateji ve liderlik ilginiz Selahaddin'e uygun"},
            {"name": "NapolÃ©on Bonaparte", "reason": "Askeri strateji merakÄ±nÄ±z NapolÃ©on'la eÅŸleÅŸiyor"}
        ])

    # Sanat ve estetik
    art_keywords = ["sanat", "gÃ¼zel", "estetik", "yaratÄ±cÄ±", "ilham", "ÅŸiir"]
    if any(keyword in all_text for keyword in art_keywords):
        suggestions.extend([
            {"name": "Michelangelo", "reason": "Sanat ve yaratÄ±cÄ±lÄ±k ilginiz Michelangelo ile uyumlu"},
            {"name": "Fuzuli", "reason": "Estetik anlayÄ±ÅŸÄ±nÄ±z Fuzuli'nin ÅŸiirine yakÄ±n"}
        ])

    # Bilim ve keÅŸif
    science_keywords = ["bilim", "keÅŸif", "araÅŸtÄ±rma", "deney", "gÃ¶zlem", "doÄŸa"]
    if any(keyword in all_text for keyword in science_keywords):
        suggestions.extend([
            {"name": "Galileo Galilei", "reason": "Bilimsel merakÄ±nÄ±z Galileo'nun ruhunu yansÄ±tÄ±yor"},
            {"name": "Ä°bn Sina", "reason": "AraÅŸtÄ±rma tutkunu Ä°bn Sina'ya Ã§ok benziyor"}
        ])

    return random.choice(suggestions) if suggestions else None


# Anasayfa butonu - Sol Ã¼st kÃ¶ÅŸe
col_home, col_title = st.columns([1, 10])
with col_home:
    if st.button("ğŸ ", help="Anasayfaya dÃ¶n"):
        st.session_state.current_page = "home"
        st.session_state.current_conversation_id = None
        st.session_state.current_character = ""
        st.session_state.messages = []
        st.session_state.time_travel_active = False
        st.rerun()

with col_title:
    st.title("ğŸ§™â€â™‚ HistorAI - Tarihi Karakter Chatbotu")

# SaÄŸ panel: Sohbet geÃ§miÅŸi
with st.sidebar:
    st.header("ğŸ“š GeÃ§miÅŸ Sohbetler")

    # Filtreleme
    filter_char = st.text_input("Karaktere gÃ¶re filtrele")

    # Sohbetleri getir (sabitlenenler Ã¶nce)
    if filter_char:
        c.execute("""SELECT id, character, title, is_pinned, conversation_type FROM conversations 
                    WHERE character LIKE ? ORDER BY is_pinned DESC, created_at DESC""",
                  ('%' + filter_char + '%',))
    else:
        c.execute("""SELECT id, character, title, is_pinned, conversation_type FROM conversations 
                    ORDER BY is_pinned DESC, created_at DESC""")
    conversations = c.fetchall()

    # Sohbet listesi
    for conv_id, char, title, is_pinned, conv_type in conversations:
        pin_icon = "ğŸ“Œ " if is_pinned else ""
        type_icon = "â° " if conv_type == "time_travel" else ""
        label = f"{pin_icon}{type_icon}{char}: {title[:20]}..."

        col1, col2, col3 = st.columns([4, 1, 1])
        with col1:
            if st.button(label, key=f"conv_{conv_id}"):
                st.session_state.current_conversation_id = conv_id
                st.session_state.current_character = char
                st.session_state.current_page = "chat"
                # Mevcut sohbetin mesajlarÄ±nÄ± yÃ¼kle
                c.execute("SELECT question, answer FROM messages WHERE conversation_id = ? ORDER BY created_at",
                          (conv_id,))
                messages = c.fetchall()
                st.session_state.messages = []
                for q, a in messages:
                    st.session_state.messages.append({"role": "user", "content": q})
                    st.session_state.messages.append({"role": "assistant", "content": a})
                st.rerun()

        with col2:
            # Pin/Unpin butonu
            pin_text = "ğŸ“Œ" if not is_pinned else "ğŸ“"
            if st.button(pin_text, key=f"pin_{conv_id}"):
                new_pin_status = 0 if is_pinned else 1
                c.execute("UPDATE conversations SET is_pinned = ? WHERE id = ?",
                          (new_pin_status, conv_id))
                conn.commit()
                st.rerun()

        with col3:
            # Sil butonu
            if st.button("ğŸ—‘", key=f"del_{conv_id}"):
                c.execute("DELETE FROM messages WHERE conversation_id = ?", (conv_id,))
                c.execute("DELETE FROM conversations WHERE id = ?", (conv_id,))
                conn.commit()
                if st.session_state.current_conversation_id == conv_id:
                    st.session_state.current_conversation_id = None
                    st.session_state.messages = []
                    st.session_state.current_page = "home"
                st.rerun()

    st.divider()

    # Yeni sohbet baÅŸlat
    if st.button("âœ¨ Yeni Sohbet BaÅŸlat"):
        st.session_state.current_conversation_id = None
        st.session_state.current_character = ""
        st.session_state.messages = []
        st.session_state.current_page = "home"
        st.session_state.time_travel_active = False
        st.rerun()

    # TÃ¼m geÃ§miÅŸi sil
    if st.button("ğŸ§¨ TÃ¼m GeÃ§miÅŸi Sil"):
        c.execute("DELETE FROM messages")
        c.execute("DELETE FROM conversations")
        conn.commit()
        st.session_state.current_conversation_id = None
        st.session_state.messages = []
        st.session_state.current_page = "home"
        st.rerun()

    st.divider()

    # Ä°ndirme seÃ§enekleri
    st.subheader("ğŸ“¥ Ä°ndirme SeÃ§enekleri")

    if st.session_state.current_conversation_id:
        st.write("Mevcut sohbeti indir:")


        # PDF indirme fonksiyonu
        def create_conversation_pdf(conversation_id):
            if not REPORTLAB_AVAILABLE:
                st.error("PDF oluÅŸturmak iÃ§in ReportLab gerekli: pip install reportlab")
                return None

            # Sohbet bilgilerini al
            c.execute("SELECT character, title FROM conversations WHERE id = ?", (conversation_id,))
            conv_info = c.fetchone()
            if not conv_info:
                return None

            character, title = conv_info
            c.execute("SELECT question, answer FROM messages WHERE conversation_id = ? ORDER BY created_at",
                      (conversation_id,))
            messages = c.fetchall()

            def register_modern_fonts():
                """Modern font kaydÄ±"""
                try:
                    if platform.system() == "Windows":
                        font_paths = [
                            "C:/Windows/Fonts/segoeui.ttf",  # Segoe UI - modern
                            "C:/Windows/Fonts/calibri.ttf",
                            "C:/Windows/Fonts/arial.ttf"
                        ]
                        bold_paths = [
                            "C:/Windows/Fonts/segoeuib.ttf",
                            "C:/Windows/Fonts/calibrib.ttf",
                            "C:/Windows/Fonts/arialbd.ttf"
                        ]
                    elif platform.system() == "Darwin":
                        font_paths = [
                            "/System/Library/Fonts/SF-Pro-Text-Regular.otf",
                            "/System/Library/Fonts/Helvetica.ttc",
                            "/Library/Fonts/Arial.ttf"
                        ]
                        bold_paths = [
                            "/System/Library/Fonts/SF-Pro-Text-Bold.otf",
                            "/System/Library/Fonts/Helvetica-Bold.ttc",
                            "/Library/Fonts/Arial Bold.ttf"
                        ]
                    else:
                        font_paths = [
                            "/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf",
                            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
                        ]
                        bold_paths = [
                            "/usr/share/fonts/truetype/ubuntu/Ubuntu-B.ttf",
                            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
                        ]

                    # Font kaydÄ±
                    for i, (regular, bold) in enumerate(zip(font_paths, bold_paths)):
                        try:
                            if os.path.exists(regular) and os.path.exists(bold):
                                pdfmetrics.registerFont(TTFont('ModernFont', regular))
                                pdfmetrics.registerFont(TTFont('ModernFont-Bold', bold))
                                return 'ModernFont'
                        except:
                            continue

                    return 'Helvetica'
                except:
                    return 'Helvetica'

            modern_font = register_modern_fonts()
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50,
                                    topMargin=50, bottomMargin=50)

            styles = getSampleStyleSheet()

            # Modern stiller
            title_style = ParagraphStyle(
                'ModernTitle',
                parent=styles['Heading1'],
                fontSize=24,
                alignment=1,
                spaceAfter=30,
                textColor=HexColor('#1a365d'),
                fontName=f'{modern_font}-Bold' if modern_font != 'Helvetica' else 'Helvetica-Bold'
            )

            subtitle_style = ParagraphStyle(
                'ModernSubtitle',
                parent=styles['Heading2'],
                fontSize=16,
                alignment=1,
                spaceAfter=20,
                textColor=HexColor('#2d3748'),
                fontName=f'{modern_font}-Bold' if modern_font != 'Helvetica' else 'Helvetica-Bold'
            )

            question_style = ParagraphStyle(
                'ModernQuestion',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=8,
                spaceBefore=15,
                textColor=HexColor('#2b6cb0'),
                fontName=f'{modern_font}-Bold' if modern_font != 'Helvetica' else 'Helvetica-Bold'
            )

            answer_style = ParagraphStyle(
                'ModernAnswer',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=15,
                leading=16,
                textColor=HexColor('#1a202c'),
                fontName=modern_font
            )

            content = []
            content.append(Paragraph("ğŸ§™â€â™‚ HistorAI Sohbeti", title_style))
            content.append(Paragraph(f"Karakter: {character}", subtitle_style))
            content.append(Spacer(1, 30))

            for i, (question, answer) in enumerate(messages, 1):
                def clean_text(text):
                    return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                content.append(Paragraph(f"<b>â“ Soru {i}:</b> {clean_text(question)}", question_style))
                content.append(Paragraph(f"<b>ğŸ’¬ {character}:</b> {clean_text(answer)}", answer_style))

                if i < len(messages):
                    content.append(Spacer(1, 10))

            try:
                doc.build(content)
                buffer.seek(0)
                return buffer
            except Exception as e:
                st.error(f"PDF oluÅŸtururken hata: {str(e)}")
                return None


        # Word dosyasÄ± oluÅŸturma
        def create_conversation_word(conversation_id):
            if not DOCX_AVAILABLE:
                st.error("Word dosyasÄ± oluÅŸturmak iÃ§in python-docx gerekli: pip install python-docx")
                return None

            c.execute("SELECT character, title FROM conversations WHERE id = ?", (conversation_id,))
            conv_info = c.fetchone()
            if not conv_info:
                return None

            character, title = conv_info
            c.execute("SELECT question, answer FROM messages WHERE conversation_id = ? ORDER BY created_at",
                      (conversation_id,))
            messages = c.fetchall()

            doc = Document()

            # BaÅŸlÄ±k
            title_para = doc.add_heading('ğŸ§™â€â™‚ HistorAI Sohbeti', 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Karakter bilgisi
            char_para = doc.add_heading(f'Karakter: {character}', level=1)
            char_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph('')

            # Mesajlar
            for i, (question, answer) in enumerate(messages, 1):
                # Soru
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f'â“ Soru {i}: ')
                q_run.bold = True
                q_run.font.color.rgb = RGBColor(43, 108, 176)
                q_para.add_run(question)

                # Cevap
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f'ğŸ’¬ {character}: ')
                a_run.bold = True
                a_run.font.color.rgb = RGBColor(212, 84, 58)
                a_para.add_run(answer)

                doc.add_paragraph('')

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer


        # JSON oluÅŸturma
        def create_conversation_json(conversation_id):
            c.execute("SELECT character, title FROM conversations WHERE id = ?", (conversation_id,))
            conv_info = c.fetchone()
            if not conv_info:
                return None

            character, title = conv_info
            c.execute("SELECT question, answer FROM messages WHERE conversation_id = ? ORDER BY created_at",
                      (conversation_id,))
            messages = c.fetchall()

            data = {
                "character": character,
                "title": title,
                "messages": [{"question": q, "answer": a} for q, a in messages]
            }

            return BytesIO(json.dumps(data, indent=4, ensure_ascii=False).encode("utf-8"))


        # Ä°ndirme butonlarÄ±
        col1, col2, col3 = st.columns(3)

        with col1:
            pdf_data = create_conversation_pdf(st.session_state.current_conversation_id)
            if pdf_data:
                st.download_button("ğŸ“„ PDF", data=pdf_data,
                                   file_name=f"historai_{st.session_state.current_character}.pdf",
                                   mime="application/pdf")

        with col2:
            word_data = create_conversation_word(st.session_state.current_conversation_id)
            if word_data:
                st.download_button("ğŸ“ Word", data=word_data,
                                   file_name=f"historai_{st.session_state.current_character}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with col3:
            json_data = create_conversation_json(st.session_state.current_conversation_id)
            if json_data:
                st.download_button("ğŸ—‚ JSON", data=json_data,
                                   file_name=f"historai_{st.session_state.current_character}.json",
                                   mime="application/json")

# KiÅŸilik testi veri yapÄ±sÄ±
PERSONALITY_TEST = {
    "questions": [
        {
            "question": "Yeni bir proje baÅŸlatÄ±rken hangi yaklaÅŸÄ±mÄ± tercih edersiniz?",
            "options": [
                {"text": "DetaylÄ± plan yapar, her aÅŸamayÄ± Ã¶nceden hesaplarÄ±m",
                 "traits": {"conscientiousness": 2, "openness": 1}},
                {"text": "Genel bir fikir ile baÅŸlar, yol boyunca ÅŸekillendiriririm",
                 "traits": {"openness": 2, "conscientiousness": -1}},
                {"text": "BaÅŸkalarÄ±nÄ±n fikirlerini dinler, ortak karar veririm",
                 "traits": {"agreeableness": 2, "extraversion": 1}},
                {"text": "Ä°Ã§gÃ¼dÃ¼lerime gÃ¼venir, spontane hareket ederim", "traits": {"neuroticism": 1, "openness": 1}}
            ]
        },
        {
            "question": "KarÅŸÄ±laÅŸtÄ±ÄŸÄ±nÄ±z zorluklar karÅŸÄ±sÄ±nda nasÄ±l tepki verirsiniz?",
            "options": [
                {"text": "Analitik dÃ¼ÅŸÃ¼nÃ¼r, sistematik Ã§Ã¶zÃ¼mler ararÄ±m",
                 "traits": {"conscientiousness": 2, "neuroticism": -1}},
                {"text": "YaratÄ±cÄ± ve sÄ±ra dÄ±ÅŸÄ± yÃ¶ntemler denerim", "traits": {"openness": 2, "conscientiousness": -1}},
                {"text": "DiÄŸer insanlardan yardÄ±m ve tavsiye alÄ±rÄ±m",
                 "traits": {"agreeableness": 2, "extraversion": 1}},
                {"text": "Duygusal yaklaÅŸÄ±r, iÃ§sel motivasyonuma gÃ¼venirim",
                 "traits": {"neuroticism": 1, "agreeableness": 1}}
            ]
        },
        {
            "question": "Ä°deal bir akÅŸam nasÄ±l geÃ§irirsiniz?",
            "options": [
                {"text": "Kitap okuyarak veya Ã¶ÄŸrendiÄŸim konularÄ± derinleÅŸtirerek",
                 "traits": {"openness": 2, "extraversion": -1}},
                {"text": "ArkadaÅŸlarÄ±mla sohbet ederek, deneyimlerimi paylaÅŸarak",
                 "traits": {"extraversion": 2, "agreeableness": 1}},
                {"text": "Sanat, mÃ¼zik veya yaratÄ±cÄ± aktivitelerle",
                 "traits": {"openness": 2, "conscientiousness": -1}},
                {"text": "DÃ¼zenli rutinlerimi sÃ¼rdÃ¼rerek, planlarÄ±mÄ± gÃ¶zden geÃ§irerek",
                 "traits": {"conscientiousness": 2, "extraversion": -1}}
            ]
        },
        {
            "question": "Liderlik tarzÄ±nÄ±zÄ± nasÄ±l tanÄ±mlarsÄ±nÄ±z?",
            "options": [
                {"text": "Vizyon sahibi, ilham verici ve yenilikÃ§i", "traits": {"openness": 2, "extraversion": 1}},
                {"text": "Disiplinli, adaletli ve kurallara baÄŸlÄ±",
                 "traits": {"conscientiousness": 2, "agreeableness": 1}},
                {"text": "Empati kuran, destekleyici ve iÅŸbirlikÃ§i", "traits": {"agreeableness": 2, "extraversion": 1}},
                {"text": "KararlÄ±, tutarlÄ± ancak esnek", "traits": {"conscientiousness": 1, "neuroticism": -1}}
            ]
        },
        {
            "question": "Hangi tÃ¼r bilgi sizi en Ã§ok cezbeder?",
            "options": [
                {"text": "Bilimsel keÅŸifler ve teknolojik yenilikler",
                 "traits": {"openness": 2, "conscientiousness": 1}},
                {"text": "Tarihsel olaylar ve kÃ¼ltÃ¼rel geliÅŸmeler", "traits": {"openness": 1, "conscientiousness": 1}},
                {"text": "Ä°nsan iliÅŸkileri ve sosyal dinamikler", "traits": {"agreeableness": 2, "extraversion": 1}},
                {"text": "Felsefe ve yaÅŸamÄ±n anlamÄ± Ã¼zerine dÃ¼ÅŸÃ¼nceler", "traits": {"openness": 2, "neuroticism": 1}}
            ]
        }
    ],
    "characters": [
        {
            "name": "Leonardo da Vinci",
            "traits": {"openness": 10, "conscientiousness": 7, "extraversion": 5, "agreeableness": 6, "neuroticism": 4},
            "description": "Ã‡ok yÃ¶nlÃ¼ deha, sanat ve bilimi birleÅŸtiren yaratÄ±cÄ± vizyon",
            "quote": "Ã–ÄŸrenme bizim yaÅŸadÄ±ÄŸÄ±mÄ±z sÃ¼rece devam eder."
        },
        {
            "name": "Fatih Sultan Mehmet",
            "traits": {"openness": 8, "conscientiousness": 9, "extraversion": 8, "agreeableness": 5, "neuroticism": 3},
            "description": "Stratejik dÃ¼ÅŸÃ¼nÃ¼r, kararlÄ± lider ve vizyon sahibi fatih",
            "quote": "Ya Ä°stanbul'u alÄ±rÄ±m, ya da Ä°stanbul beni alÄ±r."
        },
        {
            "name": "Mevlana",
            "traits": {"openness": 9, "conscientiousness": 6, "extraversion": 6, "agreeableness": 10, "neuroticism": 2},
            "description": "Sevgi dolu, hoÅŸgÃ¶rÃ¼lÃ¼ ve hakikati arayan mutasavvÄ±f",
            "quote": "Sevgi yolculuÄŸu, bizi kendimize gÃ¶tÃ¼rÃ¼r."
        },
        {
            "name": "Ibn Khaldun",
            "traits": {"openness": 9, "conscientiousness": 8, "extraversion": 4, "agreeableness": 7, "neuroticism": 3},
            "description": "Sosyal bilimcÄ±, tarihÃ§i ve medeniyet analisti",
            "quote": "Tarih, toplumlarÄ±n yÃ¼kseliÅŸ ve Ã§Ã¶kÃ¼ÅŸ kanunlarÄ±nÄ± Ã¶ÄŸretir."
        },
        {
            "name": "Yunus Emre",
            "traits": {"openness": 8, "conscientiousness": 5, "extraversion": 7, "agreeableness": 9, "neuroticism": 4},
            "description": "HalkÄ±n ozanÄ±, sevgi ve kardeÅŸlik ÅŸairi",
            "quote": "YaratÄ±lanÄ± severiz, Yaratan'dan Ã¶tÃ¼rÃ¼."
        },
        {
            "name": "Sultan Alparslan",
            "traits": {"openness": 6, "conscientiousness": 9, "extraversion": 7, "agreeableness": 6, "neuroticism": 2},
            "description": "Adil hÃ¼kÃ¼mdar, stratejik komutan ve devlet adamÄ±",
            "quote": "Adalet, saltanatÄ±n temelidir."
        },
        {
            "name": "Ä°bn Sina (Avicenna)",
            "traits": {"openness": 10, "conscientiousness": 8, "extraversion": 4, "agreeableness": 7, "neuroticism": 3},
            "description": "Hekim, filozof ve bilim insanÄ±",
            "quote": "Bilgi, onu arayan ve emek verenlerindir."
        },
        {
            "name": "AkÅŸemseddin",
            "traits": {"openness": 8, "conscientiousness": 8, "extraversion": 5, "agreeableness": 8, "neuroticism": 2},
            "description": "Bilgin, mutasavvÄ±f ve Fatih'in hocasÄ±",
            "quote": "Ä°lim Ã¶ÄŸren, kendini bil."
        }
    ]
}


# Test sonucu hesaplama fonksiyonu
def calculate_personality_match(user_scores):
    """KullanÄ±cÄ±nÄ±n kiÅŸilik puanlarÄ±nÄ± karakterlerle eÅŸleÅŸtir"""
    best_matches = []

    for character in PERSONALITY_TEST["characters"]:
        # Her karakter iÃ§in benzerlik puanÄ± hesapla
        similarity_score = 0
        total_possible = 0

        for trait in ["openness", "conscientiousness", "extraversion", "agreeableness", "neuroticism"]:
            user_score = user_scores.get(trait, 0)
            char_score = character["traits"][trait]

            # Mutlak farkÄ± hesapla (0-10 arasÄ±)
            diff = abs(user_score - char_score)
            similarity = 10 - diff  # Fark ne kadar az ise benzerlik o kadar yÃ¼ksek

            similarity_score += similarity
            total_possible += 10

        # YÃ¼zde olarak hesapla
        match_percentage = (similarity_score / total_possible) * 100

        best_matches.append({
            "character": character,
            "percentage": match_percentage
        })

    # En yÃ¼ksek eÅŸleÅŸenleri dÃ¶ndÃ¼r
    return sorted(best_matches, key=lambda x: x["percentage"], reverse=True)


# Session state iÃ§in test deÄŸiÅŸkenleri
if "test_active" not in st.session_state:
    st.session_state.test_active = False
if "test_question_index" not in st.session_state:
    st.session_state.test_question_index = 0
if "test_scores" not in st.session_state:
    st.session_state.test_scores = {"openness": 0, "conscientiousness": 0, "extraversion": 0, "agreeableness": 0,
                                    "neuroticism": 0}
if "test_completed" not in st.session_state:
    st.session_state.test_completed = False

# Ana iÃ§erik - Sayfa yÃ¶nlendirmesi
if st.session_state.current_page == "home" and not st.session_state.current_conversation_id:
    # ANASAYFA

    # Karakter Tavsiyesi Motoru - EÄŸer Ã¶nceki sohbetler varsa
    if len(st.session_state.messages) >= 4:
        suggestion = analyze_conversation_style(st.session_state.messages)
        if suggestion:
            st.markdown("---")
            st.markdown("### ğŸ¯ Size Ã–zel Karakter Tavsiyesi")
            col1, col2 = st.columns([3, 1])
            with col1:
                st.info(f"**{suggestion['name']}** - {suggestion['reason']}")
            with col2:
                if st.button(f"ğŸ’¬ {suggestion['name']} ile sohbet et", key="suggestion_chat"):
                    st.session_state.current_character = suggestion['name']
                    c.execute("INSERT INTO conversations (character, title, conversation_type) VALUES (?, ?, ?)",
                              (suggestion['name'], f"{suggestion['name']} ile tavsiye sohbeti", "normal"))
                    st.session_state.current_conversation_id = c.lastrowid
                    conn.commit()
                    st.session_state.current_page = "chat"
                    st.rerun()

    # Zamanda Yolculuk BÃ¶lÃ¼mÃ¼
    st.markdown("---")
    st.markdown("### â° Zamanda Yolculuk - Olay AnÄ± CanlandÄ±rma")
    st.markdown("*Tarihin en kritik anlarÄ±na gidip o dÃ¶nemin karakterleriyle yaÅŸayÄ±n!*")

    if not st.session_state.time_travel_active:
        col1, col2 = st.columns([2, 1])
        with col1:
            selected_year = st.selectbox(
                "Hangi tarihi olayÄ±n ortasÄ±na gitmek istersiniz?",
                options=list(HISTORICAL_EVENTS.keys()),
                format_func=lambda x: f"{x} - {HISTORICAL_EVENTS[x]['event']} ({HISTORICAL_EVENTS[x]['date']})"
            )

            if st.button("ğŸš€ Zamanda YolculuÄŸa BaÅŸla", type="primary"):
                st.session_state.time_travel_active = True
                st.session_state.selected_event = selected_year
                st.rerun()

        with col2:
            st.markdown("#### ğŸ­ Deneyim:")
            st.markdown(
                "ğŸ¬ Sinematik giriÅŸ  \nğŸ‘¥ Otomatik karakter eÅŸleÅŸmesi  \nğŸŒ Ortam betimlemesi  \nğŸ¯ Interaktif roleplay")

    else:
        # Zamanda yolculuk aktif
        event_data = HISTORICAL_EVENTS[st.session_state.selected_event]

        st.markdown(f"### ğŸŒ {event_data['event']} - {event_data['date']}")

        # Sinematik giriÅŸ
        st.markdown(f"""
        <div style="
            background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%);
            padding: 25px;
            border-radius: 15px;
            color: white;
            margin: 20px 0;
            border-left: 5px solid #f39c12;
        ">
            <h3>ğŸ¬ Zamanda Yolculuk BaÅŸlÄ±yor...</h3>
            <p style="font-size: 16px; line-height: 1.6;">
                <strong>Ortam:</strong> {event_data['setting']}
            </p>
            <p style="font-size: 18px; font-style: italic; margin-top: 20px;">
                "{event_data['opening']}"
            </p>
        </div>
        """, unsafe_allow_html=True)

        # Karakter seÃ§imi
        st.markdown("#### ğŸ‘¥ Bu olayda kiminle karÅŸÄ±laÅŸmak istersiniz?")
        selected_character = st.radio(
            "Karakter seÃ§in:",
            event_data['characters'],
            horizontal=True
        )

        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("ğŸ­ Bu Karakterle Sohbete BaÅŸla", type="primary"):
                st.session_state.current_character = selected_character
                # Zamanda yolculuk sohbeti oluÅŸtur
                c.execute("INSERT INTO conversations (character, title, conversation_type) VALUES (?, ?, ?)",
                          (selected_character, f"{event_data['event']} - {selected_character}", "time_travel"))
                st.session_state.current_conversation_id = c.lastrowid
                conn.commit()

                # Ä°lk mesajÄ± otomatik olarak ekle
                opening_message = f"Zamanda yolculuk yaparak {event_data['date']} tarihindeki {event_data['event']} olayÄ±nÄ±n tam ortasÄ±ndayÄ±m. {event_data['opening']}"
                st.session_state.messages = [{"role": "user", "content": opening_message}]

                st.session_state.time_travel_active = False
                st.session_state.current_page = "chat"
                st.rerun()

        with col1:
            if st.button("â†© Geri DÃ¶n"):
                st.session_state.time_travel_active = False
                st.rerun()

    # TarihÃ® KiÅŸilik Testi BÃ¶lÃ¼mÃ¼
    st.markdown("---")

    col1, col2 = st.columns([2, 1])

    with col1:
        st.markdown("### ğŸ§¬ TarihÃ® KiÅŸilik Testi")
        st.markdown("*5 soruluk kÄ±sa testle hangi tarihi karaktere benzediÄŸinizi keÅŸfedin!*")

        if not st.session_state.test_active and not st.session_state.test_completed:
            if st.button("ğŸš€ Teste BaÅŸla", type="primary"):
                st.session_state.test_active = True
                st.session_state.test_question_index = 0
                st.session_state.test_scores = {"openness": 0, "conscientiousness": 0, "extraversion": 0,
                                                "agreeableness": 0, "neuroticism": 0}
                st.rerun()

    with col2:
        st.markdown("#### ğŸ¯ Test SonrasÄ±:")
        st.markdown("âœ¨ KiÅŸilik eÅŸleÅŸmesi  \nğŸ“Š Uyumluluk yÃ¼zdesi  \nğŸ’¬ Direkt sohbet baÅŸlat")

    # Test aktifse sorularÄ± gÃ¶ster
    if st.session_state.test_active:
        current_q = st.session_state.test_question_index
        total_q = len(PERSONALITY_TEST["questions"])

        if current_q < total_q:
            st.markdown("---")
            # Progress bar
            progress = (current_q) / total_q
            st.progress(progress, text=f"Soru {current_q + 1} / {total_q}")

            question_data = PERSONALITY_TEST["questions"][current_q]

            st.markdown(f"### ğŸ“ Soru {current_q + 1}")
            st.markdown(f"{question_data['question']}")

            # SeÃ§enekleri radio button olarak gÃ¶ster
            option_labels = [opt["text"] for opt in question_data["options"]]

            selected_option = st.radio(
                "SeÃ§iminizi yapÄ±n:",
                options=range(len(option_labels)),
                format_func=lambda x: option_labels[x],
                key=f"test_q_{current_q}"
            )

            col1, col2, col3 = st.columns([1, 1, 2])

            with col2:
                if st.button("â¡ Sonraki Soru", type="primary"):
                    # SeÃ§ilen seÃ§eneÄŸin trait puanlarÄ±nÄ± ekle
                    selected_traits = question_data["options"][selected_option]["traits"]
                    for trait, score in selected_traits.items():
                        st.session_state.test_scores[trait] += score

                    st.session_state.test_question_index += 1
                    st.rerun()

            with col1:
                if st.button("âŒ Testi Durdur"):
                    st.session_state.test_active = False
                    st.session_state.test_question_index = 0
                    st.rerun()

        else:
            # Test tamamlandÄ± - sonuÃ§larÄ± gÃ¶ster
            st.session_state.test_active = False
            st.session_state.test_completed = True
            st.rerun()

    # Test sonuÃ§larÄ±
    if st.session_state.test_completed:
        st.markdown("---")
        st.markdown("## ğŸ‰ Test SonuÃ§larÄ±nÄ±z")

        # SkorlarÄ± normalize et (0-10 arasÄ±)
        normalized_scores = {}
        for trait, score in st.session_state.test_scores.items():
            # Her soru iÃ§in maksimum 2 puan alÄ±nabilir, 5 soru var
            max_possible = 10  # 5 soru x 2 puan
            min_possible = -5  # 5 soru x -1 puan (bazÄ± negatif skorlar var)

            # 0-10 arasÄ±na normalize et
            normalized = ((score - min_possible) / (max_possible - min_possible)) * 10
            normalized_scores[trait] = max(0, min(10, normalized))

        # En iyi eÅŸleÅŸmeleri bul
        matches = calculate_personality_match(normalized_scores)

        # En iyi 3 eÅŸleÅŸmeyi gÃ¶ster
        for i, match in enumerate(matches[:3]):
            character = match["character"]
            percentage = match["percentage"]

            if i == 0:
                # En iyi eÅŸleÅŸme - Ã¶zel stil
                st.markdown(f"""
                <div style="
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    padding: 20px;
                    border-radius: 15px;
                    color: white;
                    margin: 10px 0;
                    text-align: center;
                ">
                    <h2>ğŸ† En Ä°yi EÅŸleÅŸmeniz!</h2>
                    <h1>{character['name']}</h1>
                    <h2>%{percentage:.0f} Uyumluluk</h2>
                    <p style="font-style: italic;">"{character['quote']}"</p>
                    <p>{character['description']}</p>
                </div>
                """, unsafe_allow_html=True)

                # Direkt sohbet baÅŸlatma butonu
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button(f"ğŸ’¬ {character['name']} ile Sohbet BaÅŸlat", type="primary", key="start_chat_best"):
                        st.session_state.current_character = character['name']
                        c.execute("INSERT INTO conversations (character, title, conversation_type) VALUES (?, ?, ?)",
                                  (character['name'], f"{character['name']} ile kiÅŸilik testi sohbeti", "personality"))
                        st.session_state.current_conversation_id = c.lastrowid
                        conn.commit()
                        st.session_state.test_completed = False
                        st.session_state.current_page = "chat"
                        st.rerun()

            else:
                # DiÄŸer eÅŸleÅŸmeler
                with st.expander(f"#{i + 1} - {character['name']} (%{percentage:.0f} uyumluluk)"):
                    st.markdown(f"{character['description']}")
                    st.markdown(f"\"{character['quote']}\"")

                    if st.button(f"ğŸ’¬ {character['name']} ile Sohbet BaÅŸlat", key=f"start_chat_{i}"):
                        st.session_state.current_character = character['name']
                        c.execute("INSERT INTO conversations (character, title, conversation_type) VALUES (?, ?, ?)",
                                  (character['name'], f"{character['name']} ile kiÅŸilik testi sohbeti", "personality"))
                        st.session_state.current_conversation_id = c.lastrowid
                        conn.commit()
                        st.session_state.test_completed = False
                        st.session_state.current_page = "chat"
                        st.rerun()

        # Testi tekrar alma butonu
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("ğŸ”„ Testi Tekrar Al"):
                st.session_state.test_active = False
                st.session_state.test_completed = False
                st.session_state.test_question_index = 0
                st.rerun()

        with col2:
            if st.button("â¡ Manuel Karakter SeÃ§"):
                st.session_state.test_completed = False
                st.rerun()

    # Manuel karakter seÃ§imi (test yapÄ±lmadÄ±ysa veya manuel seÃ§im isteniyorsa)
    if not st.session_state.test_active and not st.session_state.test_completed:
        st.markdown("---")
        st.markdown("### ğŸ­ Veya Manuel Karakter SeÃ§in")
        character = st.text_input("Tarihi karakter adÄ±nÄ± girin:",
                                  placeholder="Ã–rn: Fatih Sultan Mehmet, Leonardo da Vinci, Mevlana...")

        if character:
            st.session_state.current_character = character
            # Yeni conversation oluÅŸtur
            c.execute("INSERT INTO conversations (character, title, conversation_type) VALUES (?, ?, ?)",
                      (character, f"{character} ile sohbet", "manual"))
            st.session_state.current_conversation_id = c.lastrowid
            conn.commit()
            st.session_state.current_page = "chat"
            st.rerun()

elif st.session_state.current_page == "chat" or st.session_state.current_conversation_id:
    # SOHBET SAYFASI
    st.markdown(f"### ğŸ—£ {st.session_state.current_character} ile sohbet ediyorsunuz")

    # Karakter tavsiyesi (sohbet sÄ±rasÄ±nda)
    if len(st.session_state.messages) >= 6:  # 3 soru-cevap dÃ¶ngÃ¼sÃ¼nden sonra
        suggestion = analyze_conversation_style(st.session_state.messages)
        if suggestion and suggestion['name'] != st.session_state.current_character:
            with st.expander("ğŸ¯ Size baÅŸka bir karakter Ã¶nerisi var!", expanded=False):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.info(f"**{suggestion['name']}** - {suggestion['reason']}")
                with col2:
                    if st.button(f"ğŸ’¬ {suggestion['name']}", key="mid_chat_suggestion"):
                        # Mevcut sohbeti kaydet ve yeni karakter ile baÅŸla
                        if st.session_state.messages:
                            first_question = st.session_state.messages[0]["content"] if st.session_state.messages[0][
                                                                                            "role"] == "user" else "Sohbet"
                            title = first_question[:50] + "..." if len(first_question) > 50 else first_question
                            c.execute("UPDATE conversations SET title = ? WHERE id = ?",
                                      (title, st.session_state.current_conversation_id))
                            conn.commit()

                        st.session_state.current_character = suggestion['name']
                        c.execute("INSERT INTO conversations (character, title, conversation_type) VALUES (?, ?, ?)",
                                  (suggestion['name'], f"{suggestion['name']} ile tavsiye sohbeti", "suggestion"))
                        st.session_state.current_conversation_id = c.lastrowid
                        conn.commit()
                        st.session_state.messages = []
                        st.rerun()

    # Sohbet geÃ§miÅŸini gÃ¶ster
    for message in st.session_state.messages:
        if message["role"] == "user":
            with st.chat_message("user"):
                st.write(message["content"])
        else:
            with st.chat_message("assistant"):
                st.write(message["content"])

    # Yeni mesaj giriÅŸi
    if prompt := st.chat_input("Sorunuzu yazÄ±n..."):
        # KullanÄ±cÄ± mesajÄ±nÄ± ekle
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)

        # AI yanÄ±tÄ± oluÅŸtur
        with st.chat_message("assistant"):
            with st.spinner("YanÄ±t oluÅŸturuluyor..."):
                # Zamanda yolculuk sohbeti iÃ§in Ã¶zel prompt
                if len(st.session_state.messages) == 1 and "Zamanda yolculuk" in st.session_state.messages[0][
                    "content"]:
                    ai_prompt = f"""
Sen {st.session_state.current_character} olarak, kullanÄ±cÄ±nÄ±n zamanda yolculuk yaparak seni ziyaret ettiÄŸi bu Ã¶zel anÄ± canlandÄ±rÄ±yorsun.

ZAMANDA YOLCULUK ROLEPLAY KURALLARI:
1. Atmosferi ve ortamÄ± detaylÄ± betimle
2. O dÃ¶nemin gerginliÄŸi, kokularÄ±, sesleri, gÃ¶rÃ¼ntÃ¼leri dahil et
3. Karakterin o anki ruh hali ve durumunu yansÄ±t
4. KullanÄ±cÄ±yÄ± bu tarihi olayÄ±n bir parÃ§asÄ± gibi hissettir
5. DÃ¶nem diline uygun ama anlaÅŸÄ±lÄ±r ÅŸekilde konuÅŸ

KullanÄ±cÄ±nÄ±n mesajÄ±: {prompt}

Bu tarihi anÄ± tam olarak yaÅŸatarak, kendini {st.session_state.current_character} olarak tanÄ±t ve durumu betimle.
"""
                else:
                    ai_prompt = f"""
Sen yalnÄ±zca tarihsel olarak belgelenmiÅŸ, gerÃ§ek ve yaÅŸamÄ±ÅŸ karakterlerin rolÃ¼nÃ¼ yapabilirsin...

TEMEL KURALLAR:
1. YalnÄ±zca insanlÄ±k tarihinde yaÅŸamÄ±ÅŸ, gÃ¼venilir tarihsel kaynaklarda yer alan kiÅŸiliklerin yerine geÃ§ebilirsin.
2. Her yanÄ±tÄ±n tarihsel olarak doÄŸrulanabilir olmalÄ±. Uydurma bilgi, tahmin ya da kurgu iÃ§erik Ã¼retmek kesinlikle yasaktÄ±r.

ROL YAPMAYI REDDETMEN GEREKEN DURUMLAR:
- GerÃ§ek olmayan, hayali veya anlamsÄ±z karakterler (Ã¶rneÄŸin: "Merhaba", "Kral Ejder", "Mehmet", "RobotX")
- Tarihsel figÃ¼r olmayan Ã§aÄŸdaÅŸ kiÅŸiler (Ã¶rneÄŸin: Elon Musk, Donald Trump, Britney Spears, Ronaldo)
- TÃ¼rkiye Cumhuriyeti tarafÄ±ndan hassas kabul edilen kiÅŸi ve iÃ§erikler (Ã¶rneÄŸin: terÃ¶r Ã¶rgÃ¼tleri ve terÃ¶r Ã¶rgÃ¼tÃ¼ kurucularÄ±, suÃ§lu nitelikteki insanlar)
- Dini, tanrÄ±sal veya kutsal figÃ¼rler (Ã¶rneÄŸin: TanrÄ±, Hz. Muhammed, Ä°sa)
- KÃ¼fÃ¼r, cinsellik, hakaret ve toplumsal olarak hassas konular

Bu tÃ¼r isteklerde:
- NazikÃ§e isteÄŸi reddet
- KÄ±sa aÃ§Ä±klama yap: "Bu kiÅŸi/talep, rol yapabileceÄŸim gÃ¼venilir tarihsel iÃ§eriklere uygun deÄŸildir."
- Kesinlikse hiÃ§bir ÅŸekilde rol yapma veya bu kiÅŸiler adÄ±na konuÅŸma.

Tarihsel Uydurma YasaÄŸÄ±:
EÄŸer kullanÄ±cÄ± sana gerÃ§ek bir tarihi olayla ilgisi olmayan bir hikÃ¢ye, konuÅŸma, anÄ± ya da deneyim sorduysa:
- Uydurma cevap verme.
- "Bu olay/kaynak tarihsel olarak doÄŸrulanmÄ±ÅŸ deÄŸildir." diyerek aÃ§Ä±klama yap.
- Ancak istenen olay gerÃ§ekte yaÅŸanmÄ±ÅŸsa, tarihsel bilgiye dayalÄ± ÅŸekilde cevap verebilirsin.

---

Åimdi {st.session_state.current_character} olarak konuÅŸuyorsun. AÅŸaÄŸÄ±daki soruyu, bu karakterin tarihsel gerÃ§eklerine ve dÃ¶nemin diline sadÄ±k kalarak cevapla:

- Soru: {prompt}
"""
                try:
                    response = model.generate_content(ai_prompt)
                    answer = response.text
                    st.write(answer)

                    # MesajlarÄ± kaydet
                    st.session_state.messages.append({"role": "assistant", "content": answer})
                    c.execute("INSERT INTO messages (conversation_id, question, answer) VALUES (?, ?, ?)",
                              (st.session_state.current_conversation_id, prompt, answer))
                    conn.commit()

                except Exception as e:
                    error_msg = f"Bir hata oluÅŸtu: {str(e)}"
                    st.error(error_msg)
                    st.session_state.messages.append({"role": "assistant", "content": error_msg})

    # Sohbeti bitir butonu
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("âœ… Sohbeti Bitir"):
            # Conversation title'Ä±nÄ± gÃ¼ncelle
            if st.session_state.messages:
                first_question = st.session_state.messages[0]["content"] if st.session_state.messages[0][
                                                                                "role"] == "user" else "Sohbet"
                title = first_question[:50] + "..." if len(first_question) > 50 else first_question
                c.execute("UPDATE conversations SET title = ? WHERE id = ?",
                          (title, st.session_state.current_conversation_id))
                conn.commit()

            st.session_state.current_conversation_id = None
            st.session_state.current_character = ""
            st.session_state.messages = []
            st.session_state.current_page = "home"
            st.success("Sohbet tamamlandÄ± ve geÃ§miÅŸe kaydedildi!")
            st.rerun()